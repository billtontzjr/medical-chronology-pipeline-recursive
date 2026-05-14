"""Output renderers: JSON + DOCX.

Produces the final-form deliverables from the assembled markdown
chronology and the verified header. The DOCX rendering matches
Dr. Tontz's standing format: Times New Roman 12pt, single-spaced,
justified, with a header block for patient identity + report date.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.assembler import HEADER_PLACEHOLDER, UNDATED_HEADING
from src.header_extractor import Header
from src.schemas import VerifiedFact


ENTRY_DATE_RE = re.compile(r"^\s*(\d{2}/\d{2}/\d{4})\.")


@dataclass
class JsonEntry:
    date: Optional[str]
    facility: Optional[str]
    provider_name: Optional[str]
    provider_credentials: Optional[str]
    visit_type: Optional[str]
    narrative: str
    supporting_fact_ids: List[str]


# ---------------------------------------------------------------- markdown io

def render_header_block(header: Header, report_date: Optional[str] = None) -> str:
    """Produce the multi-line header that goes at the top of chronology.md."""
    report_date = report_date or datetime.now().strftime("%m/%d/%Y")
    name = (header.patient_name or "PATIENT NAME UNKNOWN").upper()
    return "\n".join(
        [
            "MEDICAL RECORDS SUMMARY",
            name,
            f"Date of Birth: {header.dob or 'Unknown'}",
            f"Date of Injury: {header.doi or 'Unknown'}",
            f"Report Date: {report_date}",
        ]
    )


def inject_header_into_chronology(
    markdown_path: Path,
    header: Header,
    report_date: Optional[str] = None,
) -> None:
    """Replace ``HEADER_PLACEHOLDER`` in ``chronology.md`` with the real header."""
    content = markdown_path.read_text(encoding="utf-8")
    block = render_header_block(header, report_date=report_date)
    if HEADER_PLACEHOLDER in content:
        content = content.replace(HEADER_PLACEHOLDER, block, 1)
    else:
        content = block + "\n\n" + content
    markdown_path.write_text(content, encoding="utf-8")


# ------------------------------------------------------------------- json

def _parse_entry_first_sentence(entry_text: str) -> dict:
    m = ENTRY_DATE_RE.match(entry_text)
    if not m:
        return dict(date=None, facility=None, provider_name=None, provider_credentials=None, visit_type=None)
    date_str = m.group(1)
    rest = entry_text[m.end():].lstrip()
    parts = rest.split(". ", 3)
    facility = parts[0].strip() if parts else None
    provider_name = None
    provider_credentials = None
    visit_type = None
    if len(parts) > 1:
        prov = parts[1].strip()
        if "," in prov:
            provider_name, _, creds = prov.partition(",")
            provider_name = provider_name.strip()
            provider_credentials = creds.strip()
        else:
            provider_name = prov
    if len(parts) > 2:
        visit_type = parts[2].strip().rstrip(".")
    return dict(
        date=date_str,
        facility=facility or None,
        provider_name=provider_name or None,
        provider_credentials=provider_credentials or None,
        visit_type=visit_type or None,
    )


def render_chronology_json(
    header: Header,
    markdown_chronology: str,
    verified_facts: List[VerifiedFact],
    *,
    output_path: Path,
    report_date: Optional[str] = None,
) -> dict:
    """Write chronology.json with metadata + entries.

    Each entry's ``supporting_fact_ids`` is a list of chunk_id strings
    of facts that share the visit_date+facility (best-effort linking;
    the assembler does not currently emit explicit fact ids).
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", markdown_chronology) if p.strip()]
    entries: List[JsonEntry] = []
    for para in paragraphs:
        if UNDATED_HEADING in para:
            continue
        first_line = para.splitlines()[0]
        if not ENTRY_DATE_RE.match(first_line):
            continue
        meta = _parse_entry_first_sentence(para)
        supporting = [
            f.chunk_id
            for f in verified_facts
            if f.visit_date == meta["date"]
            and (
                meta["facility"] is None
                or (f.facility or "").lower() in (meta["facility"] or "").lower()
                or (meta["facility"] or "").lower() in (f.facility or "").lower()
            )
        ]
        entries.append(
            JsonEntry(
                date=meta["date"],
                facility=meta["facility"],
                provider_name=meta["provider_name"],
                provider_credentials=meta["provider_credentials"],
                visit_type=meta["visit_type"],
                narrative=para,
                supporting_fact_ids=sorted(set(supporting)),
            )
        )

    payload = {
        "metadata": {
            "patient_name": header.patient_name,
            "date_of_birth": header.dob,
            "date_of_injury": header.doi,
            "report_date": report_date or datetime.now().strftime("%m/%d/%Y"),
            "generated": datetime.now().isoformat(timespec="seconds"),
            "verified_fact_count": len(verified_facts),
            "entry_count": len(entries),
        },
        "entries": [e.__dict__ for e in entries],
    }
    output_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return payload


# ----------------------------------------------------------------- docx

def _set_default_font(doc: Document) -> None:
    for style_name in ("Normal", "Default Paragraph Font"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)


def _add_justified_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)


def render_chronology_docx(
    markdown_chronology: str,
    header: Header,
    output_path: Path,
    *,
    report_date: Optional[str] = None,
) -> None:
    """Render the markdown chronology as a DOCX file."""
    doc = Document()
    _set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(72)

    # Header block
    block = render_header_block(header, report_date=report_date).splitlines()
    for i, line in enumerate(block):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        if i == 0:
            p.runs[0].bold = False  # spec says no bolding even for header

    doc.add_paragraph("")  # blank separator

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", markdown_chronology) if p.strip()]
    for para in paragraphs:
        if para.startswith(HEADER_PLACEHOLDER) or para.startswith("MEDICAL RECORDS SUMMARY"):
            continue
        if UNDATED_HEADING in para:
            doc.add_paragraph("")
            p = doc.add_paragraph("Undated Entries")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        clean = para.replace("—", ",")  # em-dash → comma (defensive)
        _add_justified_paragraph(doc, clean)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


__all__ = [
    "render_header_block",
    "inject_header_into_chronology",
    "render_chronology_json",
    "render_chronology_docx",
    "JsonEntry",
]
